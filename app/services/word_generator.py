"""
Word document generation module.
"""
import logging
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.config import settings

logger = logging.getLogger(__name__)


class WordGenerator:
    """Generates Word documents from formatted text."""
    
    def __init__(self):
        self.output_dir = Path(settings.OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    async def create_document(self, client_name: str, content: str, output_subdir: str = None, date_range: str = None) -> str:
        """
        Create a Word document for a client's conversations.
        
        Args:
            client_name: Name of the client
            content: Formatted conversation text
            output_subdir: Optional override output directory (e.g., "output-2")
            date_range: Optional date range string to display in title (e.g., "November 1-11, 2025")
            
        Returns:
            Path to the generated file
        """
        try:
            # Sanitize client name for filename
            safe_client_name = self._sanitize_filename(client_name)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_client_name}_{timestamp}.docx"
            out_dir = self.output_dir if not output_subdir else Path(output_subdir)
            out_dir.mkdir(parents=True, exist_ok=True)
            file_path = out_dir / filename
            
            # Create document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add title with date range if provided
            logger.info(f"Creating document for {client_name}, date_range={date_range}")
            if date_range:
                title_text = f'Conversations with {client_name} ({date_range})'
            else:
                title_text = f'Conversations with {client_name}'
            logger.info(f"Document title: {title_text}")
            title = doc.add_heading(title_text, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date
            date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para_format = date_para.runs[0].font
            date_para_format.size = Pt(10)
            date_para_format.italic = True
            
            # Add spacing
            doc.add_paragraph()
            
            # Add content
            # Split content by lines and format appropriately
            lines = content.split('\n')
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    # Empty line - add spacing
                    doc.add_paragraph()
                    current_paragraph = None
                    continue
                
                # Check if line is a header (starts with =)
                if line.startswith('='):
                    # Add as heading
                    doc.add_heading(line.replace('=', '').strip(), level=2)
                    current_paragraph = None
                # Check if line is a conversation line (format: Speaker Name: text [time])
                elif ':' in line and not line.startswith('='):
                    # Parse: Speaker Name: text [00:00 - 00:05] or Speaker Name: text
                    if '[' in line and ']' in line:
                        # Has timestamp at the end
                        # Split to separate text and timestamp
                        last_bracket = line.rfind(']')
                        text_part = line[:last_bracket + 1].rstrip()
                        # Extract timestamp (everything in brackets)
                        timestamp_start = text_part.rfind('[')
                        if timestamp_start != -1:
                            timestamp = text_part[timestamp_start:]
                            text_with_speaker = text_part[:timestamp_start].strip()
                        else:
                            timestamp = ""
                            text_with_speaker = text_part
                    else:
                        # No timestamp
                        timestamp = ""
                        text_with_speaker = line
                    
                    # Split speaker and text
                    if ':' in text_with_speaker:
                        speaker, text = text_with_speaker.split(':', 1)
                        current_paragraph = doc.add_paragraph()
                        
                        # Add speaker name in bold
                        run1 = current_paragraph.add_run(speaker.strip() + ':')
                        run1.font.bold = True
                        run1.font.size = Pt(11)
                        
                        # Add text
                        if text.strip():
                            current_paragraph.add_run(' ' + text.strip())
                        
                        # Add timestamp in smaller italic font at the end
                        if timestamp:
                            current_paragraph.add_run(' ')
                            run2 = current_paragraph.add_run(timestamp)
                            run2.font.size = Pt(9)
                            run2.font.italic = True
                            run2.font.color.rgb = None  # Use default gray-ish color
                else:
                    # Regular text line
                    if current_paragraph:
                        current_paragraph.add_run(' ' + line)
                    else:
                        doc.add_paragraph(line)
            
            # Save document
            doc.save(str(file_path))
            logger.info(f"Generated Word document: {file_path}")
            
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Error creating Word document for {client_name}: {str(e)}")
            raise
    
    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitize filename by removing invalid characters.
        
        Args:
            filename: Original filename
            
        Returns:
            Sanitized filename
        """
        # Remove or replace invalid filename characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # Remove leading/trailing spaces and dots
        filename = filename.strip('. ')
        
        # Limit length
        if len(filename) > 100:
            filename = filename[:100]
        
        return filename

